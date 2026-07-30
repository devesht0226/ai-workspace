"""Research agent: RAG + knowledge graph + memory synthesis."""

from __future__ import annotations

from io import BytesIO

import httpx
from bs4 import BeautifulSoup
from sqlalchemy import select
from sqlalchemy.orm import Session

from app.core.config import get_settings
from app.core.exceptions import ValidationAppError
from app.models import DocumentChunk, User
from app.providers.llm import ChatMessage
from app.services import documents as document_service
from app.services import knowledge_graph as kg
from app.services import memory as memory_service


def search_web(query: str, *, limit: int = 5) -> list[dict[str, str]]:
    """Return public DuckDuckGo HTML results without requiring an API key."""
    settings = get_settings()
    if settings.environment == "test" or settings.llm_provider == "fake":
        return []
    try:
        with httpx.Client(timeout=10.0, follow_redirects=True) as client:
            response = client.get(
                "https://html.duckduckgo.com/html/",
                params={"q": query},
                headers={"User-Agent": "AI-Workspace/1.0"},
            )
            response.raise_for_status()
        soup = BeautifulSoup(response.text, "html.parser")
        results = []
        for item in soup.select(".result"):
            title = item.select_one(".result__a")
            snippet = item.select_one(".result__snippet")
            if title:
                results.append(
                    {
                        "title": title.get_text(" ", strip=True),
                        "url": str(title.get("href") or ""),
                        "snippet": snippet.get_text(" ", strip=True) if snippet else "",
                    }
                )
            if len(results) >= limit:
                break
        return results
    except (httpx.HTTPError, ValueError):
        return []


def _empty_rag(message: str = "No document evidence available.") -> dict:
    return {"answer": message, "citations": []}


def _compact_graph(graph: dict) -> dict:
    matches = graph.get("matches") or graph.get("nodes") or []
    if isinstance(matches, list):
        return {"matches": matches[:8]}
    return {"matches": []}


def _build_brief(
    question: str,
    *,
    rag_answer: str,
    citations: list,
    graph: dict,
    mem_ctx: str,
    web_results: list[dict[str, str]],
) -> str:
    lines = [
        "# Research brief",
        f"**Question:** {question}",
        "",
        "## Document evidence (RAG)",
        rag_answer.strip() or "No relevant document passages found.",
    ]
    if citations:
        lines.append("")
        lines.append("### Citations")
        for cite in citations[:5]:
            if isinstance(cite, dict):
                name = cite.get("filename") or "document"
                page = cite.get("page_number")
                snippet = str(cite.get("snippet") or "")[:180]
            else:
                name = getattr(cite, "filename", "document")
                page = getattr(cite, "page_number", None)
                snippet = str(getattr(cite, "snippet", "") or "")[:180]
            page_bit = f" p.{page}" if page is not None else ""
            lines.append(f"- {name}{page_bit}: {snippet}")

    lines.extend(["", "## Knowledge graph"])
    matches = (graph or {}).get("matches") or []
    if matches:
        for match in matches[:8]:
            lines.append(f"- {match}")
    else:
        lines.append("- No strong graph matches for this query.")

    lines.extend(["", "## Memory"])
    lines.append(mem_ctx.strip() if mem_ctx else "- No prior memory for this topic.")

    lines.extend(["", "## Web leads (unverified)"])
    if web_results:
        for item in web_results[:5]:
            title = item.get("title") or "Result"
            url = item.get("url") or ""
            snippet = item.get("snippet") or ""
            lines.append(f"- {title}: {snippet[:160]} ({url})")
    else:
        lines.append("- No web results (offline, blocked, or empty).")

    # Honest note for demo questions with no SLA docs
    lowered = question.lower()
    if "sla" in lowered and not citations:
        lines.extend(
            [
                "",
                "## Note",
                "No SLA document is indexed yet. Upload an availability/SLA PDF under Documents, "
                "then re-run research for grounded answers.",
            ]
        )
    return "\n".join(lines)


def run_research(db: Session, user: User, question: str, *, model_family: str | None = None) -> dict:
    rag_payload = _empty_rag()
    try:
        rag = document_service.query_documents(db, user, question=question, top_k=4)
        rag_payload = {
            "answer": rag.answer,
            "citations": [c.model_dump(mode="json") for c in rag.citations],
        }
    except Exception as exc:  # noqa: BLE001
        rag_payload = _empty_rag(f"RAG step skipped: {exc}")

    try:
        graph = kg.query_graph(db, user, question.split()[0] if question.split() else question)
    except Exception:  # noqa: BLE001
        graph = {"matches": []}
    graph = _compact_graph(graph)

    try:
        mem_ctx = memory_service.memory_context(db, user, question) or ""
    except Exception:  # noqa: BLE001
        mem_ctx = ""

    web_results = search_web(question)

    settings = get_settings()
    family = model_family or "llama"
    synthesis = _build_brief(
        question,
        rag_answer=str(rag_payload.get("answer") or ""),
        citations=list(rag_payload.get("citations") or []),
        graph=graph,
        mem_ctx=mem_ctx,
        web_results=web_results,
    )

    # Optional short LLM polish — never fail the whole research run if it errors
    if settings.environment != "test" and settings.llm_provider != "fake":
        try:
            from app.providers.router import get_model_router

            routed = get_model_router().chat(
                [
                    ChatMessage(
                        role="user",
                        content=(
                            "Polish this research brief into 6-10 clear sentences. "
                            "Keep citations and the distinction between document evidence "
                            "vs unverified web leads. Do not invent facts.\n\n"
                            f"{synthesis[:3500]}"
                        ),
                    )
                ],
                family=model_family,
            )
            polished = (routed.get("content") or "").strip()
            if polished and len(polished) > 40:
                synthesis = polished
            family = routed.get("family") or family
        except Exception:  # noqa: BLE001
            # Keep structured brief
            pass
    else:
        family = model_family or "llama"

    try:
        memory_service.remember(
            db,
            user,
            content=f"Research on '{question}': {synthesis[:400]}",
            memory_type="long_term",
            source="research_agent",
            importance=2,
        )
    except Exception:  # noqa: BLE001
        pass

    return {
        "agent": "research",
        "model_family": family,
        "brief": synthesis,
        "rag": rag_payload,
        "knowledge_graph": graph,
        "memory_used": bool(mem_ctx),
        "web_results": web_results,
    }


def compare_documents(db: Session, user: User, doc_a_id, doc_b_id) -> dict:
    """Compare the leading document chunks while enforcing document ownership."""
    doc_a = document_service.get_document(db, user, doc_a_id)
    doc_b = document_service.get_document(db, user, doc_b_id)

    def first_chunks(document_id) -> str:
        chunks = list(
            db.scalars(
                select(DocumentChunk)
                .where(DocumentChunk.document_id == document_id)
                .order_by(DocumentChunk.chunk_index)
                .limit(3)
            ).all()
        )
        return "\n\n".join(chunk.content for chunk in chunks)[:6000]

    text_a, text_b = first_chunks(doc_a.id), first_chunks(doc_b.id)
    if not text_a or not text_b:
        raise ValidationAppError("Both documents need indexed content before comparison")
    settings = get_settings()
    if settings.environment == "test" or settings.llm_provider == "fake":
        comparison = (
            f"Comparison of {doc_a.filename} and {doc_b.filename}.\n"
            f"Document A preview: {text_a[:400]}\nDocument B preview: {text_b[:400]}"
        )
        family = "llama"
    else:
        from app.providers.router import get_model_router

        routed = get_model_router().chat(
            [
                ChatMessage(
                    role="user",
                    content=(
                        "Compare these two documents. Identify shared themes, key differences, "
                        "and any contradictions. Be concise.\n\n"
                        f"DOCUMENT A ({doc_a.filename}):\n{text_a}\n\n"
                        f"DOCUMENT B ({doc_b.filename}):\n{text_b}"
                    ),
                )
            ]
        )
        comparison, family = routed["content"], routed["family"]
    return {
        "document_a": {"id": str(doc_a.id), "filename": doc_a.filename},
        "document_b": {"id": str(doc_b.id), "filename": doc_b.filename},
        "comparison": comparison,
        "model_family": family,
    }


def export_brief_markdown(brief: str) -> str:
    return f"# Research Brief\n\n{brief.strip()}\n"


def export_brief_docx(brief: str) -> bytes:
    from docx import Document as DocxDocument

    document = DocxDocument()
    document.add_heading("Research Brief", level=1)
    for paragraph in brief.strip().splitlines():
        document.add_paragraph(paragraph)
    output = BytesIO()
    document.save(output)
    return output.getvalue()
